import os
from docx import Document

def remove_headers(doc):
    for section in doc.sections:
        # 处理三种可能的页眉类型
        for header in [section.header, section.even_page_header, section.first_page_header]:
            if header.is_linked_to_previous:
                continue
            # 清除所有段落内容
            for paragraph in header.paragraphs:
                paragraph.clear()
            # 补充删除页眉中的表格（如果存在）
            for table in header.tables:
                try:
                    table._element.getparent().remove(table._element)
                except:
                    pass

def process_directory(root_dir):
    if not os.path.exists(root_dir):  # 新增目录存在性检查
        print(f"目录不存在: {root_dir}")
        return
    # 在遍历前添加权限检查
    if not os.access(root_dir, os.W_OK):
        print(f"目录不可写: {root_dir}")
        return

    for root, dirs, files in os.walk(root_dir):
        # 修改为更健壮的异常处理结构
        try:
            for file in files:
                if not file.endswith('.docx'):  # 添加格式过滤
                    continue
                path = os.path.join(root, file)
                print(f"正在处理: {path}")  # 添加进度提示
                try:
                    doc = Document(path)
                    remove_headers(doc)
                    doc.save(path)
                except Exception as e:
                    # 增强错误分类提示
                    error_msg = str(e).lower()
                    if 'cannot open file' in error_msg:
                        print("→ 文件可能已被其他程序打开")
                    elif 'password' in error_msg or 'encrypted' in error_msg:
                        print("→ 文档受密码保护")
                    else:
                        print(f"→ 未知错误: {type(e).__name__}")
                    continue
        except Exception as e:
            print(f"遍历过程中发生错误: {type(e).__name__}")
            continue

# 使用示例
if __name__ == '__main__':
    docs_dir = os.path.abspath('docs')
    process_directory(docs_dir)
