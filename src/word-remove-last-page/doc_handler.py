from docx import Document
import os
import sys
import xml.etree.ElementTree as ET

def remove_last_page(input_path, output_path):
    try:
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"输入文件不存在: {input_path}")

        doc = Document(input_path)

        # 使用更可靠的分页检测方式
        delete_paragraphs = []
        page_breaks = [
            i for i, p in enumerate(doc.paragraphs)
            if 'w:br' in p._element.xml and 'type="page"' in p._element.xml
        ]

        # 当存在分页符时，保留最后一页前的所有内容
        if page_breaks:
            last_break = page_breaks[-1]
            delete_paragraphs = doc.paragraphs[last_break+1:]
        else:
            # 无分页符时删除最后三个段落
            delete_paragraphs = doc.paragraphs[-3:] if len(doc.paragraphs) > 3 else []

        # 逆向删除避免索引错位
        for p in reversed(delete_paragraphs):
            try:
                p._element.getparent().remove(p._element)
            except AttributeError:
                continue

        doc.save(output_path)

    except ET.ParseError as e:
        print(f"XML解析错误: {str(e)}")
        sys.exit(2)
    except Exception as e:
        print(f"未知错误: {str(e)}")
        sys.exit(3)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python doc_handler.py <输入文件> <输出文件>")
        sys.exit(1)
    remove_last_page(sys.argv[1], sys.argv[2])
